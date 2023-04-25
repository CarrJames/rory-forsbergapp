from flask import Flask, render_template, request, redirect, url_for, session, send_file
from flask_wtf import FlaskForm
from flask_wtf.file import FileField, FileRequired, FileAllowed
from wtforms import StringField, SubmitField
from wtforms.validators import DataRequired, Email
from flask_uploads import UploadSet, configure_uploads
import pandas as pd
from flask_sqlalchemy import SQLAlchemy
from docx import Document
from docx.shared import Inches
import os, requests, shutil, json, googlemaps, smtplib, ssl, folium
import geopy.distance
from datetime import datetime
from PIL import Image
from email.message import EmailMessage
import geopandas as gpd
from rtree import index
from shapely.geometry import Point
from docx.enum.table import WD_TABLE_ALIGNMENT

app = Flask(__name__)
app.config['SECRET_KEY'] = 'secret_key'
app.config['UPLOADED_CSV_DEST'] = 'uploads/csv'
#   DATABASE CONFIGURATION
app.config['SQLALCHEMY_DATABASE_URI'] = 'sqlite:///userlogs.db'
app.config['SQLALCHEMY_TRACK_MODIFICATIONS'] = False
db = SQLAlchemy(app)
# counter for pano webpage output (solves cache problem)
counter = 0
#   FILE UPLOAD CONFIGURATION
csv_uploads = UploadSet('csv')
configure_uploads(app, csv_uploads)
# Spatial Index configuration (doesnt work inside the cell-tower function)
idx = index.Index()
column_names = ['latitude', 'longitude']
df = pd.read_csv(r'celltowers\234-revised.csv', names=column_names, header=None)
# Converting it to a geopandas df
gdf = gpd.GeoDataFrame(df, geometry=gpd.points_from_xy(df.longitude, df.latitude))
# filling the spatial index with the geometry bounds
for i, tower in gdf.iterrows():
    if tower.geometry is not None:
        idx.insert(i, tower.geometry.bounds)

# creating the database model
class User(db.Model):
    __tablename__ = 'userlogs'
    id = db.Column(db.Integer, primary_key=True)
    username = db.Column(db.String(20))
    email = db.Column(db.String(100))
    csv_file = db.Column(db.String(100), nullable=False)
    added = db.Column(db.DateTime, default=datetime.now)

    def __init__(self, username, email, csv_file):
        self.username = username
        self.email = email
        self.csv_file = csv_file
# call to create the database model
def init_db():
    db.drop_all()
    db.create_all()
    new_user = User(username='user1', email='test@test.com', csv_file='test.csv')
    db.session.add(new_user)    
    db.session.commit()
# deleting all prior entries to the database
@app.route('/clear_db', methods=['POST'])
def clear_db():
    all_users = db.session.query(User).all()
    for user in all_users:
        db.session.delete(user)
    db.session.commit()
    return render_template('logs.html')
# form set up for the index page
class MyForm(FlaskForm):
    username = StringField('Username', validators=[DataRequired()])
    email = StringField('Email', validators=[DataRequired(), Email()])
    csv_file = FileField('CSV File', validators=[
        FileRequired(),
        FileAllowed(['csv'],'FILE FORMATE MUST BE .CSV')
    ])
    submit = SubmitField('Submit')
@app.route('/', methods=['GET', 'POST'])
def index():
    empty_folders()
    delete_pano_templates()
    session['logged_in'] = False
    form = MyForm()
    if form.validate_on_submit():
        with open('templates/pano.html', 'w'):
            pass
        username = form.username.data
        email = form.email.data
        csv_filename = csv_uploads.save(form.csv_file.data)
        # Adding the inputs to the database
        new_user = User(username=form.username.data, email=form.email.data, csv_file=csv_filename)
        db.session.add(new_user)
        db.session.commit()   
        # adding the email and csv to the session so I can access them in different functions
        session['email'] = email
        session['csv_filename'] = csv_filename
        #checking columns are formatted correctly
        locations = pd.read_csv('uploads/csv/' + csv_filename)
        expected_columns = ['latitude (deg)', 'longitude (deg)']
        # check if the DataFrame has all the expected columns
        if set(expected_columns).issubset(set(locations.columns)):
            print('inputted csv has correct columns')
            formatted_address()
            session['logged_in'] = True
            return redirect(url_for('success'))
        else:
            print('inputted csv has inccorect columns')
            return render_template('error.html')
        
    return render_template('index.html', form=form)
# returning full page of pano images
@app.route('/formatted')
def formatted():
    # fixing issue with pano.html not reloading
    return render_template(session.get('pano'), timestamp=datetime.now())
@app.route('/success')
def success():
    file_size = os.path.getsize('output.docx')
    print(file_size)
    # checking if the output document is larger than the allowed emailing file size
    if file_size > 25000000:
        print('file size over 25mb')
        return render_template('success.html', button_disabled=True)
    else:
        print('file size under 25mb')
        return render_template('success.html', button_disabled=False)
# shows database entries
@app.route('/logs')
def logs():
    data=User.query.all()
    return render_template('logs.html', data=data)
# func for showing closest cell towers
@app.route('/celldist')
def celldist():
    get_closest_towers()
    coords_filename = f'uploads/csv/{session.get("csv_filename")}'
    coords_df = pd.read_csv(coords_filename)
    results = pd.read_csv('results.csv')
    # mapping it using folium 
    m = folium.Map(location=[52.7, -1.4], zoom_start=6)
    fg1 = folium.FeatureGroup(name='Cell Towers', show=True)
    popup_str1 = f"Distance to closest point: {results.iloc[i]['distance']} miles"
    for i, row in results.iterrows():
        folium.Marker(location=[row['latitude'], row['longitude']],
                    tooltip=f"ID: {i}",
                    popup=popup_str1,
                    icon=folium.Icon(color='red')
                    ).add_to(fg1)
    
# Create a feature group for the second set of markers (green color)
    fg2 = folium.FeatureGroup(name='Inputted', show=True)
    for i, row in coords_df.iterrows():
        popup_str = f"Distance to closest tower: {results.iloc[i]['distance']} miles"
        folium.Marker(location=[row['latitude (deg)'], row['longitude (deg)']], 
                    tooltip=f"ID: {i}",
                    popup=popup_str,
                    icon=folium.Icon(color='green')
                    ).add_to(fg2)

# Add the feature groups to the map object
    fg1.add_to(m)
    fg2.add_to(m)
    folium.LayerControl().add_to(m)
    map_html = m._repr_html_()
    return render_template('celltowers.html', map=map_html)
# email function using temporary email 
@app.route('/send_email', methods=['POST'])
def send_email():
    email_sender = 'disformatter@gmail.com'
    email_password = 'bqcwnzqvxhcryylq'
    email_reciever = session.get('email')

    subject = 'Document of Formatted Addresses'
    body = 'Please find the formatted addresses document attached.'

    em = EmailMessage()

    em['From'] = email_sender
    em['To'] = email_reciever
    em['Subject'] = subject
    em.set_content(body)

    with open('output.docx', 'rb') as f:
        file_data = f.read()
    em.add_attachment(file_data, maintype='application', subtype='vnd.openxmlformats-officedocument.wordprocessingml.document', filename='output.docx')
    context = ssl.create_default_context()

    with smtplib.SMTP_SSL('smtp.gmail.com', 465, context=context) as smtp:
        smtp.login(email_sender, email_password)
        smtp.sendmail(email_sender, email_reciever, em.as_string())
    return render_template('sent.html')
@app.route('/download', methods=['POST'])
def download():
    return send_file('output.docx', as_attachment=True)
# empties the uploads folder and the static folder to reduce application size    
@app.before_first_request
def empty_folders():
    upload_folder = app.config['UPLOAD_FOLDER'] = os.path.join(os.getcwd(), 'uploads', 'csv')
    temp_folder = app.config['TEMP_FOLDER'] = os.path.join(os.getcwd(), 'static')
    out_folder = app.config['OUT_FOLDER'] = os.path.join(os.getcwd(), 'outputs')
    s_images = app.config['S_IMAGES'] = os.path.join(os.getcwd(), 'standalone_images')
    for folder in [upload_folder, temp_folder, out_folder, s_images]:
        for filename in os.listdir(folder):
            file_path = os.path.join(folder, filename)
            try:
                if os.path.isfile(file_path) or os.path.islink(file_path):
                    os.unlink(file_path)
            except Exception as e:
                print('Failed to delete %s. Reason: %s' % (file_path, e))
# address formatter using googles reverse geocoder api 
def formatted_address():#
    csv_filename = session.get('csv_filename')
    locations = pd.read_csv('uploads/csv/' + csv_filename)
    gmaps = googlemaps.Client(key='AIzaSyBwP_5ZGFGEhgo1Zc9cxW5l2jjEz5-gd1o')

    approx_address = []
    # putting the lat and long to the api
    for i in range(0,len(locations)):
        loc1 = locations[['latitude (deg)', 'longitude (deg)']].iloc[i]
        reverse_geocode_result = gmaps.reverse_geocode((loc1['latitude (deg)'], loc1['longitude (deg)']))
        jsonString = json.dumps(reverse_geocode_result)
        jsonFile = open("data.json", "w")
        jsonFile.write(jsonString)
        jsonFile.close()
        with open('data.json', 'r', encoding='utf-8') as f:
            my_data = json.load(f)
            f_address = my_data[0]['formatted_address']
            approx_address.append(f_address)

    locations['approx-address'] = approx_address
    print(csv_filename)
    locations.to_csv('outputs/'+ csv_filename, index=False)
    pano()
# pano stitching using googles static api 
def pano():
    global counter
    counter += 1
    csv_filename = session.get('csv_filename')
    locations_pano = pd.read_csv('outputs/' + csv_filename)
    print(locations_pano)
    headings = [0, 90, 180, 270]
    locations_pano['img_source'] = locations_pano['approx-address']
    for i in range(0, len(locations_pano)):
        lat1 = locations_pano.iloc[i]['latitude (deg)']
        long1 = locations_pano.iloc[i]['longitude (deg)']
        address = locations_pano.iloc[i]['approx-address']
        for heading in headings:
            input_heading = heading
            url = "https://maps.googleapis.com/maps/api/streetview?location={},{}&size=640x640&pitch=0&fov=90&heading={}&key=AIzaSyBwP_5ZGFGEhgo1Zc9cxW5l2jjEz5-gd1o".format(lat1, long1, input_heading)
            response = requests.get(url,stream=True)
            if response.status_code == 200:
                with open(f"standalone_images/{heading}.jpg", "wb") as f:
                    response.raw.decode_content = True
                    shutil.copyfileobj(response.raw, f)

    ## CREATING PANO IMAGE BY STITCHING
        filenames = ["standalone_images/0.jpg", "standalone_images/90.jpg", "standalone_images/180.jpg", "standalone_images/270.jpg"]

        images = [Image.open(filename) for filename in filenames]
        widths, heights = zip(*(i.size for i in images))

        total_width = sum(widths)
        max_height = max(heights)

        pano_image = Image.new("RGB", (total_width, max_height), color="white")

        x_offset = 0
        for x, image in enumerate(images):
            pano_image.paste(image, (x_offset, 0))
            x_offset += widths[x]
        ## OUTPUTTING PANO IMAGES WITH ADDRESS NAME AS FILE NAME
        filename = str(i) +'.jpg'
        pano_image.save('static/' + filename)
        width = 1280
        height = 320
        locations_pano['img_source'].iloc[i] = r"<img src='{{ url_for('static', filename=" + repr(filename) + ") }}'>"
    staticmaps()
    to_word(csv_filename)
    # TO HTML
    # dropping possible columns
    expected_columns = ['GPS week', 'GPS second', 'solution status', 'height (m)']
        # check if the DataFrame has all the expected columns
    if set(expected_columns).issubset(set(locations_pano.columns)):
        locations_pano = locations_pano.drop(columns=['GPS week', 'GPS second', 'solution status', 'height (m)'])
    # putting the images together into a html file
    result_html = locations_pano.to_html()
    result_html_replaced = result_html.replace('&lt;','<').replace('&gt;','>')
    file_name = f"pano{counter}.html"
    file_path = os.path.join(os.getcwd(), 'templates', file_name)
    text_file = open(file_path, "w")
    #print(file_name, file_path)
    text_file.write('{% extends "base.html" %}')
    text_file.write('{% block content %}')
    text_file.write(result_html_replaced)
    text_file.write('{% endblock %}')
    text_file.close()
    session['pano'] = file_name
# func for finding the closest cell towers
def find_closest_towers(coords_df, gdf, idx):
    closest_towers = []
    for i, row in coords_df.iterrows():
        point = Point(row['longitude (deg)'], row['latitude (deg)'])
        # Get indices of candidate towers from spatial index
        candidate_indices = list(idx.nearest(point.bounds))
        # Calculate distances to candidate towers and find closest one
        closest_distance = float('inf')
        closest_tower = None
        for i in candidate_indices:
            tower = gdf.iloc[i]
            distance = point.distance(tower.geometry)
            if distance < closest_distance:
                closest_distance = distance
                closest_tower = tower
        closest_towers.append(closest_tower)
    return closest_towers
# gets the distance of the closest towers
def get_closest_towers():
    csv_filename = session.get('csv_filename')
    coords_df = pd.read_csv(csv_filename)
    result_df = find_closest_towers(coords_df, gdf, idx)
    results = pd.DataFrame(result_df, columns=['latitude', 'longitude', 'geometry'])
    # finding the distance of each
    for i in range(0,len(coords_df)):
        loclat = coords_df.iloc[i]['latitude (deg)']
        loclong = coords_df.iloc[i]['longitude (deg)']
        celllat = results.iloc[i]['latitude']
        celllong = results.iloc[i]['longitude']
        print('im the fucking error')
        new_column = geopy.distance.geodesic((loclat, loclong),(celllat, celllong)).miles
        results['distance'].iloc[i] = str(new_column)
        print(new_column)
    results.to_csv('results.csv')
# formatting document in a word format
def to_word(csv_filename):
    # adding distance to each cell tower
    cell_results = pd.read_csv('results.csv')
    locations_toword = pd.read_csv('outputs/' + csv_filename)
     # dropping possible columns
    expected_columns = ['GPS week', 'GPS second', 'solution status', 'height (m)']
        # check if the DataFrame has all the expected columns
    if set(expected_columns).issubset(set(locations_toword.columns)):
        locations_toword = locations_toword.drop(columns=['GPS week', 'GPS second', 'solution status', 'height (m)'])
    doc = Document()
    table = doc.add_table(rows=1, cols=7)
    table.autofit = False
    table.width = Inches(6.5)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = 'Latitude (Deg)'
    hdr_cells[1].text = 'Longitude (Deg)'
    hdr_cells[2].text = 'Location'
    hdr_cells[3].text = 'Pano Image'
    hdr_cells[4].text = 'Distance'
    hdr_cells[5].text = 'Hyperlink'
    hdr_cells[6].text = 'Map Image'

    for i in range(0,len(locations_toword)):
        lat1 = locations_toword.iloc[i]['latitude (deg)']
        long1 = locations_toword.iloc[i]['longitude (deg)']
        url = "https://www.google.com/maps/place/{},{}".format(lat1,long1)
        row_cells = table.add_row().cells
        row_cells[0].text = str(locations_toword.iloc[i]['latitude (deg)'])
        row_cells[1].text = str(locations_toword.iloc[i]['longitude (deg)'])
        row_cells[2].text = str(locations_toword.iloc[i]['approx-address'])
        cell = row_cells[3]
        paragraph = cell.add_paragraph()
        paragraph.add_run().add_picture('static/' + str(i) + '.jpg', width=Inches(2))

        row_cells[4].text = str(cell_results.iloc[i]['distance'])
        row_cells[5].text = url
        cell = row_cells[6]
        paragraph = cell.add_paragraph()
        paragraph.add_run().add_picture('maps/' + str(i) + '.jpg', width=Inches(2))

    doc.save('output.docx')
# needed as cache is saving the pano pages
def delete_pano_templates():
    template_dir = os.path.join(os.getcwd(), 'templates')
    files = os.listdir(template_dir)
    for file_name in files:
        if 'pano' in file_name:
            file_path = os.path.join(template_dir, file_name)
            os.remove(file_path)
# gives hyperlink for google map locations
def hyperlink():
    hyperlink_list = []
    locations = f'uploads/csv/{session.get("csv_filename")}'
    for i in range(0, len(locations)):
        lat1 = locations.iloc[i]['latitude (deg)']
        long1 = locations.iloc[i]['longitude (deg)']
        url = "https://www.google.com/maps/place/{},{}".format(lat1,long1)
        hyperlink_list.append(url)
    return hyperlink_list
# gives static map images of locations
def staticmaps():
    locations_name = f'uploads/csv/{session.get("csv_filename")}'
    locations = pd.read_csv(locations_name)
    for i in range(0, len(locations)):
        lat1 = locations.iloc[i]['latitude (deg)']
        long1 = locations.iloc[i]['longitude (deg)']
        url = "https://maps.googleapis.com/maps/api/staticmap?center={},{}&size=640x640&zoom=15&&markers=color:red|{},{}|&key=AIzaSyBwP_5ZGFGEhgo1Zc9cxW5l2jjEz5-gd1o".format(lat1, long1, lat1, long1)
        response = requests.get(url,stream=True)
        if response.status_code == 200:
            with open(f"maps\{i}.jpg", "wb") as f:
                response.raw.decode_content = True
                shutil.copyfileobj(response.raw, f)
# output so the user can have csv of closest cell-towers
@app.route('/output_cell_towers', methods=['POST'])
def output_cell_towers():
    coords_filename = f'uploads/csv/{session.get("csv_filename")}'
    coords_df = pd.read_csv(coords_filename)
    expected_columns = ['GPS week', 'GPS second', 'solution status', 'height (m)']
        # check if the DataFrame has all the expected columns
    if set(expected_columns).issubset(set(coords_df.columns)):
        coords_df = coords_df.drop(columns=['GPS week', 'GPS second', 'solution status', 'height (m)'])
    cell_results = pd.read_csv('results.csv')
    combined_results = pd.concat([coords_df, cell_results], axis=1)
    combined_results.columns = ['Latitude', 'Longitude', 'Tower Index', ' Tower Latitude', ' Tower Longitude', ' Tower Geometry', 'distance']
    combined_results.to_csv('combined_results.csv')
    return send_file('combined_results.csv', as_attachment=True)
@app.route('/example')
def example():
    return render_template('example.html')
@app.errorhandler(500)
def server_error(error):
    return render_template('500.html'), 500
@app.errorhandler(403)
def server_error(error):
    return render_template('403.html'), 403
@app.errorhandler(404)
def server_error(error):
    return render_template('404.html'), 404
if __name__ == '__main__':
    app.run(debug=True)
